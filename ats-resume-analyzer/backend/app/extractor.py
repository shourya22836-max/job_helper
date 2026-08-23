"""Extract plain text from PDF, DOCX, or TXT files.

PDFs with little extractable text (e.g. scanned/image-only) are automatically
fallback-OCR'd using Tesseract via pytesseract. Requires the `tesseract` binary
on the host (`brew install tesseract` on macOS, `apt install tesseract-ocr` on Debian).
"""
from __future__ import annotations

import io
import logging
import shutil
import subprocess
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".txt"}

# Heuristic: if a PDF page yields fewer than this many characters, treat it as image-only
OCR_TRIGGER_CHARS_PER_PAGE = 100


def _tesseract_available() -> bool:
    return shutil.which("tesseract") is not None


def extract_text(file_path: str | Path) -> str:
    """Return the plain text content of the given file."""
    p = Path(file_path)
    ext = p.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        return _extract_pdf(p)
    if ext == ".docx":
        return _extract_docx(p)
    return _extract_txt(p)


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            try:
                txt = page.extract_text() or ""
            except Exception as exc:  # noqa: BLE001
                logger.warning("PDF page extract failed: %s", exc)
                txt = ""
            if txt:
                parts.append(txt)
    return "\n".join(parts).strip()


def _ocr_pdf(path: Path) -> str:
    """Render each page to an image and run Tesseract OCR."""
    import pdfplumber
    from PIL import Image
    import pytesseract

    logger.info("Falling back to OCR for PDF: %s", path.name)
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            try:
                img = page.to_image(resolution=300).original
            except Exception as exc:  # noqa: BLE001
                logger.warning("Render-to-image failed on page %d: %s", idx, exc)
                continue
            try:
                txt = pytesseract.image_to_string(img)
            except Exception as exc:  # noqa: BLE001
                logger.warning("Tesseract failed on page %d: %s", idx, exc)
                txt = ""
            if txt and txt.strip():
                parts.append(txt.strip())
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_txt(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding).strip()
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def extract_text_with_meta(file_path: str | Path) -> Tuple[str, dict]:
    """Extract text and return metadata useful for ATS checks (e.g. PDF tables).

    For PDFs with little extractable text, automatically attempts OCR fallback
    if Tesseract is available on the system PATH.
    """
    p = Path(file_path)
    ext = p.suffix.lower()
    meta: dict = {
        "ext": ext,
        "pages": 0,
        "has_tables": False,
        "text_density": 0.0,
        "ocr_used": False,
        "ocr_available": _tesseract_available(),
    }

    if ext == ".pdf":
        import pdfplumber

        parts: list[str] = []
        table_count = 0
        with pdfplumber.open(p) as pdf:
            meta["pages"] = len(pdf.pages)
            for page in pdf.pages:
                try:
                    txt = page.extract_text() or ""
                except Exception:  # noqa: BLE001
                    txt = ""
                parts.append(txt)
                try:
                    if page.extract_tables():
                        table_count += 1
                except Exception:  # noqa: BLE001
                    pass
        meta["has_tables"] = table_count > 0
        full_text = "\n".join(parts).strip()
        density = len(full_text) / max(meta["pages"], 1)
        meta["text_density"] = density

        # OCR fallback: low text density suggests scanned/image-only PDF
        if (
            meta["pages"] > 0
            and density < OCR_TRIGGER_CHARS_PER_PAGE
            and meta["ocr_available"]
        ):
            ocr_text = _ocr_pdf(p)
            if ocr_text:
                full_text = ocr_text
                meta["ocr_used"] = True
                meta["text_density"] = len(full_text) / max(meta["pages"], 1)

        return full_text, meta

    text = extract_text(p)
    meta["text_density"] = len(text)
    return text, meta